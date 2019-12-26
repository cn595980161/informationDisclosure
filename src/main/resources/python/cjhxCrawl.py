import difflib
import json
import os

import requests
import time
import xlrd
import xlwt
import xlsxwriter
from docx import Document
from lxml import etree
from win32com import client as wc
from tqdm import tqdm

# import textract

url = 'http://www.cjhxfund.com/servlet/json'

domain = 'http://www.cjhxfund.com'

version = ''

year_report_path = r'.\year_report'

notice_path = r'.\notice'

headers = {
    'Content-Length': '83',
    'Content-Type': 'application/x-www-form-urlencoded; charset=UTF-8'
}


def get_fund():
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/76.0.3809.100 Safari/537.36'
    }
    r = requests.get('http://www.cjhxfund.com/main/qxcp/index.html', headers=headers)
    html = etree.HTML(r.content.decode("utf-8"))
    html_data = html.xpath('//a[@class="fundhide"]')
    list = []
    for data in html_data:
        print(data.attrib['id'] + ':' + data.text)
        if data.attrib['id'] != '001909':
            data = {'fund_code': data.attrib['id'], 'fund_name': data.text}
            list.append(data)

    return list


def crawl_notice(fund_code, catalogId, numPerPage):
    body = 'funcNo=904001&catalogId={}&numPerPage={}&isPage=Y&curpage=1&state=3&fundid=' + fund_code
    # body = funcNo=904001&catalogId=1053&numPerPage=10&isPage=Y&curpage=1&state=3&fundid=001662
    r = requests.post(url, data=body.format(catalogId, numPerPage), headers=headers)
    # print(r.status_code)
    # print(r.text)
    if r.status_code == 200:
        return json.loads(r.text)
    else:
        return None


# 获取公告
def get_fund_notice(fund_code, catalogId):
    # print(fund_code)
    result = crawl_notice(fund_code, catalogId, '1')
    # print(result['error_no'])
    if result['error_no'] == '0':
        totalRows = result['results'][0]['totalRows']
        # print(totalRows)
        result = crawl_notice(fund_code, catalogId, totalRows)
        # print(result['results'][0]['data'])
        return result['results'][0]['data']
    else:
        print('未查询到数据')
        return None


# 获取年报
def get_fund_year_report(notice_list, fund_code):
    for notice in notice_list:
        # if notice['title'].find('年年度报告') >= 0:
        if notice['title'].endswith('年年度报告'):
            # print(notice)
            download(domain + notice['link_url'], notice['title'], fund_code)


# 下载文件
def download(url, fileName, fund_code):
    res = requests.get(url)
    dir_path = year_report_path + r'\{}\{}'.format(version, fund_code)
    path = dir_path + r'\{}.{}'.format(fileName, url.split('.')[-1])
    if not os.path.exists(dir_path):
        os.makedirs(dir_path)
    with open(path, "wb") as f:
        f.write(res.content)

    # # doc转docx
    # if os.path.splitext(path)[1] == '.doc':
    #     # 转换格式
    #     doc_to_docx(path)


# 将 .doc 文件转成 .docx
def doc_to_docx(path):
    print('開始轉換格式')

    # os.system('taskkill /im wps.exe')

    path = os.path.join(os.getcwd(), path)

    w = wc.Dispatch('Kwps.Application')
    # # w = wc.Dispatch('Word.Application')
    w.Visible = 0
    w.DisplayAlerts = 0
    doc = w.Documents.Open(path)
    newpath = os.path.splitext(path)[0] + '.docx'
    doc.SaveAs(newpath, 12, False, "", True, "", False, False, False, False)
    doc.Close()
    w.Quit()

    os.remove(path)
    print('結束轉換格式')
    return newpath


# 将 .doc 文件转成 .docx
def doc_to_docx_all():
    print('开始格式转换')

    w = wc.Dispatch('Kwps.Application')
    # # w = wc.Dispatch('Word.Application')
    w.Visible = 0
    w.DisplayAlerts = 0

    rootdir = year_report_path + r'\{}'.format(version)
    if os.path.exists(rootdir):
        fund_dir = os.listdir(rootdir)
        for i in range(0, len(fund_dir)):
            fund_path = os.path.join(rootdir, fund_dir[i])
            print(fund_path)
            report_dir = os.listdir(fund_path)
            for j in range(0, len(report_dir)):
                path = os.path.join(os.getcwd(), fund_path, report_dir[j])
                print(path)

                if os.path.splitext(path)[1] == '.doc':
                    # doc转换docx
                    doc = w.Documents.Open(path)
                    newpath = os.path.splitext(path)[0] + '.docx'
                    doc.SaveAs(newpath, 12, False, "", True, "", False, False, False, False)
                    doc.Close()
                    os.remove(path)

    w.Quit()
    print('结束格式转换')


# 将 .doc 文件转成 .docx
def doc_to_docx_report(fund_code):
    print('开始格式转换')

    w = wc.Dispatch('Kwps.Application')
    # # w = wc.Dispatch('Word.Application')
    w.Visible = 0
    w.DisplayAlerts = 0

    rootdir = year_report_path + r'\{}\{}'.format(version, fund_code)
    if os.path.exists(rootdir):
        file_list = os.listdir(rootdir)
        for i in range(0, len(file_list)):
            path = os.path.join(os.getcwd(), rootdir, file_list[i])
            print(path)

            if os.path.splitext(path)[1] == '.doc':
                # doc转换docx
                doc = w.Documents.Open(path)
                newpath = os.path.splitext(path)[0] + '.docx'
                doc.SaveAs(newpath, 12, False, "", True, "", False, False, False, False)
                doc.Close()
                os.remove(path)

    w.Quit()
    print('结束格式转换')


# 读取年报
def read_year_report(path, fund_code):
    list = []

    doc = Document(path)
    year_table = read_spec_table(doc, '11.8 其他重大事件')
    print('***' + fund_code)
    for i in range(len(year_table.rows)):
        if i > 0:
            # for j in range(len(year_table.columns)):
            # print(year_table.cell(i, 1).text)
            # print(year_table.cell(i, 3).text)
            data = {'title': year_table.cell(i, 1).text, 'date': year_table.cell(i, 3).text, 'fund_code': fund_code}
            list.append(data)
    return list


# 读取年报
def read_year_reports(fund_code):
    list = []
    rootdir = year_report_path + r'\{}\{}'.format(version, fund_code)
    if os.path.exists(rootdir):
        file_list = os.listdir(rootdir)
        for i in range(0, len(file_list)):
            path = os.path.join(rootdir, file_list[i])
            if os.path.isfile(path):
                year_report = read_year_report(path, fund_code)
                list.extend(year_report)
                # print(path)

    return list


def read_spec_table(document, specText):
    paragraphs = document.paragraphs
    allTables = document.tables
    specText = specText.encode('utf-8').decode('utf-8')
    for aPara in paragraphs:
        # if aPara.text.startswith(specText):
        if aPara.text == specText:
            ele = aPara._p.getnext()
            while (ele.tag != '' and ele.tag[-3:] != 'tbl'):
                ele = ele.getnext()
            if ele.tag != '':
                for aTable in allTables:
                    if aTable._tbl == ele:
                        # for i in range(len(aTable.rows)):
                        #     for j in range(len(aTable.columns)):
                        #         print(aTable.cell(i, j).text)

                        return aTable

    return None


def create_excel(fund_code, fund_name, datas):
    # 新建一个Excel文件（只能通过新建写入）
    data = xlwt.Workbook()
    # 新建一个工作表
    table = data.add_sheet(fund_code, cell_overwrite_ok=True)
    # 写入数据到A1单元格
    # 初始化样式
    style = xlwt.XFStyle()
    style1 = xlwt.XFStyle()
    borders = xlwt.Borders()
    borders.left = 1
    borders.right = 1
    borders.top = 1
    borders.bottom = 1
    # 为样式创建字体
    font = xlwt.Font()
    # 指定字体名字
    font.name = 'Times New Roman'
    # 字体加粗
    font.bold = True
    # 将该font设定为style的字体
    style.font = font
    style.borders = borders
    style1.borders = borders
    # table.write(5, 0, u'Python Excel操作之xlwt创建表格', style)
    tabletitle = ['基金代码', '基金简称', '公告标题', '年度', '披露日期', '公告文件名']

    for i in range(0, len(tabletitle)):
        table.write(0, i, tabletitle[i], style)

    for i in range(0, len(datas)):
        # for j in range(0, len(tableA)):
        table.write(i + 1, 0, fund_code, style1)
        table.write(i + 1, 1, fund_name, style1)
        table.write(i + 1, 2, datas[i]['title'], style1)
        table.write(i + 1, 3, datas[i]['publish_date'].split('-')[0], style1)
        table.write(i + 1, 4, datas[i]['publish_date'].split(' ')[0], style1)
        table.write(i + 1, 5, datas[i]['title'], style1)

    # 注意：如果对同一个单元格重复操作，会引发overwrite Exception，想要取消该功能，需要在添加工作表时指定为可覆盖，像下面这样
    # table=data.add_sheet('name',cell_overwrite_ok=True)
    # 保存文件
    root_path = os.path.join(notice_path, str(version))
    if not os.path.exists(root_path):
        os.makedirs(root_path)
    data.save(root_path + '\\' + fund_code + '.xls')
    # 这里只能保存扩展名为xls的，xlsx的格式不支持


def create_diff_excel(datas):
    # 新建一个Excel文件
    workbook = xlsxwriter.Workbook('公告年报不一致.xlsx')
    # 新建一个工作表
    worksheet = workbook.add_worksheet()

    default_format = workbook.add_format()
    red_format = workbook.add_format({'font_color': 'red'})
    green_format = workbook.add_format({'font_color': 'green'})
    title_format = workbook.add_format({'bold': True})

    # table.write(5, 0, u'Python Excel操作之xlwt创建表格', style)
    tabletitle = ['基金代码', '基金简称', '年报公告标题', '年报披露日期', '公告标题', '披露日期', '差异']

    for i in range(0, len(tabletitle)):
        worksheet.write(0, i, tabletitle[i], title_format)

    for i in range(0, len(datas)):
        # for j in range(0, len(tableA)):
        worksheet.write(i + 1, 0, datas[i]['fund_code'])
        worksheet.write(i + 1, 1, '')
        worksheet.write(i + 1, 2, datas[i]['title'])
        worksheet.write(i + 1, 3, datas[i]['date'])
        worksheet.write(i + 1, 4, datas[i]['notic_title'])
        worksheet.write(i + 1, 5, datas[i]['notic_date'], red_format if (datas[i]['date_diff']) else default_format)
        segments = []
        for word in datas[i]['title_diff']:
            if word.startswith("+"):
                segments.append(green_format)
            elif word.startswith("-"):
                segments.append(red_format)
            else:
                segments.append(default_format)
            segments.append(word[-1])
        if len(segments) == 0:
            worksheet.write_rich_string(i + 1, 6, '')
        else:
            d = datas[i]['title_diff']
            print(d)
            worksheet.write_rich_string(i + 1, 6, *segments)

    workbook.close()


def read_excel(fund_code):
    x1 = xlrd.open_workbook(notice_path + '\{}\{}.xls'.format(version, fund_code))
    # 2、获取sheet对象
    sheet = x1.sheet_by_name(fund_code)  # 通过sheet名查找
    nrows = sheet.nrows
    list = []
    for i in range(1, nrows):
        # row = sheet.row_values(i)
        # print('第%d行值' % i, row)
        data = {'title': sheet.cell(i, 2).value, 'date': sheet.cell(i, 4).value, 'fund_code': fund_code}
        # print(data)
        list.append(data)
    return list


def string_similar(s1, s2):
    return difflib.SequenceMatcher(lambda x: x in [0, 1, 2, 3, 4, 5, 6, 7, 8, 9], s1, s2).ratio()


def compare_aaa(target_list, src_list):
    _list = []
    for target_data in target_list:
        flag = False
        date_diff = False
        similar_ratio = float()
        tmp_src_data = {}
        for src_data in src_list:

            # 一致
            if target_data['title'] == src_data['title'] and target_data['date'] == src_data['date']:
                flag = True
                break

            # 披露时间不一致
            if target_data['title'] == src_data['title'] and target_data['date'][:4] == src_data['date'][:4] and target_data['date'][5:] != src_data['date'][5:]:
                print('披露时间不一致情况:')
                data = target_data
                data['notic_title'] = src_data['title']
                data['notic_date'] = src_data['date']
                data['date_diff'] = True
                data['title_diff'] = []
                _list.append(data)
                break

            # 相似度比较
            _similar_ratio = string_similar(target_data['title'], src_data['title'])
            if similar_ratio < _similar_ratio < 1.0 and _similar_ratio >= 0.9 and target_data['date'][:4] == src_data['date'][:4]:
                similar_ratio = _similar_ratio
                tmp_src_data = src_data

        else:
            if flag is False:
                print('不一致' + str(target_data))
                data = target_data
                if tmp_src_data == {}:
                    data['notic_title'] = ''
                    data['notic_date'] = ''
                    data['date_diff'] = ''
                    data['title_diff'] = []
                else:
                    differences = list(difflib.Differ().compare(target_data['title'], tmp_src_data['title']))
                    print('不一致情况:', str(similar_ratio), "".join(differences))

                    data['title_diff'] = differences
                    data['notic_title'] = tmp_src_data['title']
                    data['notic_date'] = tmp_src_data['date']
                    data['date_diff'] = date_diff
                    data['similar_ratio'] = similar_ratio
                _list.append(data)
            continue

    return _list


def check_fund_1(fund_code, fund_name):
    # 获取公告列表
    notice_list = get_fund_notice(fund_code, '1052')
    # 获取法律列表
    law_list = get_fund_notice(fund_code, '1053')
    # 获取年报
    get_fund_year_report(notice_list, fund_code)
    # 导出excel
    create_excel(fund_code, fund_name, notice_list + law_list)


def check_fund_2(fund_code):
    # 读取年报列表
    year_report_list = read_year_reports(fund_code)
    # 读取公告列表
    notice_list = read_excel(fund_code)
    # 比对列表
    diff_list = compare_aaa(year_report_list, notice_list)
    return diff_list


if __name__ == '__main__':
    version = int(time.time())
    # version = '1577092778'
    print('当前版本号:', version)

    diff_list = []
    fund_list = get_fund()

    pbar = tqdm(fund_list)
    for fund in pbar:
        check_fund_1(fund['fund_code'], fund['fund_name'])
        pbar.set_description("进度 %s" % fund)

    doc_to_docx_all()

    pbar = tqdm(fund_list)
    for fund in pbar:
        diff = check_fund_2(fund['fund_code'])
        diff_list.extend(diff)
        pbar.set_description("进度 %s" % fund)

    print(diff_list)
    create_diff_excel(diff_list)
    # a = string_similar('创1金合信鑫收益灵活配置混合型证券投资基金招募说明书（更新）摘要（2017年第2号）', '创金合信鑫收益灵活配置混合型证券投资基金招募说明书（更新）摘要（2017年第2号）')
    # print('sdsad:' + float(a))

    # print('2018-01-01'[0:4])
